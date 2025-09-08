import pandas as pd,datetime,glob,os
import win32com.client as win32
from docx import Document
from pandas import ExcelWriter
from docx.enum.style import WD_STYLE_TYPE
from PyPDF2 import PdfMerger
#os.chdir(r"I:\Amban\Back office\בקרת השקעות- מוסדיים+מפעליות\רבעוני\פורום חוב")
def RTLWordAndExportPdf(PathFile):
    WordApp=win32.gencache.EnsureDispatch('Word.Application')
    Page = WordApp.Documents.Open(PathFile) # opens "Test" file
    WordApp.Selection.WholeStory()
    WordApp.Selection.RtlPara()
    Page.Save()
    Page.ExportAsFixedFormat(OutputFileName=PathFile[:-4] + "pdf",ExportFormat=17)
    WordApp.Quit()
    return PathFile[:-4] + "pdf"
def CreateFolder(Path,NameFolder):
    if not os.path.exists(Path+"//"+NameFolder):
        os.makedirs(Path+"//"+NameFolder)
    return Path+"//"+NameFolder
def combine_word_documents(files,Folder):
    file_ref = open("DataToPDF/Template.docx","rb")
    combined_document = Document(file_ref)
    count, number_of_files = 0, len(files)
    for file in files:
        sub_doc = Document(file)
        # Don't add a page break if you've
        # reached the last file.
        if count < number_of_files - 1:
            sub_doc.add_page_break()
        for element in sub_doc.paragraphs:
            combined_document.add_paragraph(element.text)
        count += 1
    combined_document.save('C:\\Users\\Nir.Ohad\\Desktop\\ibi-pdf\\combined_word_documents.docx')
    PathFilePdf=RTLWordAndExportPdf('C:\\Users\\Nir.Ohad\\Desktop\\ibi-pdf\\combined_word_documents.docx')
    return PathFilePdf
def get_col_widths(dataframe):
    # First we find the maximum length of the index column
    idx_max = max([len(str(s)) for s in dataframe.index.values] + [len(str(dataframe.index.name))])
    # Then, we concatenate this to the max of the lengths of column name and its values for each column, left to right
    return [idx_max] + [max([len(str(s)) for s in dataframe[col].values] + [len(col)]) for col in dataframe.columns]

#%% Create Word With all Manpik Note From benny(Bond Analsys)
#os.chdir(os.path.normpath(os.getcwd() + os.sep + os.pardir))
#MainFolder=os.path.abspath(os.getcwd())
GetManpikQNow=pd.read_excel(r"DataToPDF/Data.xlsx").rename({"מספר תיק":"MisparTik","סיווג פורום חוב":"Status"},axis=1)
GetManpikQBefore=pd.read_excel(r"DataToPDF/DataOld.xlsx").rename({"מספר תיק":"MisparTik","סיווג פורום חוב":"Status"},axis=1)
GetNameGroup=pd.read_csv(r"DataToPDF/NumberTik.csv",engine='python',encoding='ISO-8859-8')
FilterAfik=pd.read_excel(r"DataToPDF/FilterAfik.xlsx")
GetManpikQNow=GetManpikQNow.merge(FilterAfik)
GetManpikQBefore=GetManpikQNow.merge(FilterAfik)
ListStatus=["השגחה מיוחדת","בפיגור","מסופק","בעדכון"]
#GetNameGroup=GetNameGroup[(GetNameGroup["NameGroup"]=='אל על')|(GetNameGroup["NameGroup"]=='תע"ש')]
NowManpik=GetManpikQNow.query("Status in @ListStatus")[["קוד מנפיק","תאור מנפיק","Status"]]
LastQManpik=GetManpikQBefore.query("Status in @ListStatus")[["קוד מנפיק","תאור מנפיק","Status"]]
#AllManpik=NowManpik.append(LastQManpik).drop_duplicates()
AllManpik = pd.concat([NowManpik, LastQManpik]).drop_duplicates()

AllManpik.to_excel(r"DataToPDF/dsa.xlsx")
#GetNameGroup=GetNameGroup[(GetNameGroup["NameGroup"]=='עמי')]
for NameGroup in GetNameGroup["NameGroup"].drop_duplicates():
    print(NameGroup)
    try:
        #Create Word With all Manpik Note From benny(Bond Analsys)
        FolderToSave='C:\\Users\\Nir.Ohad\\Desktop\\ibi-pdf\\TempFiles'#os.path.abspath(os.getcwd())+"\\"+"קובץ מאוחד למעקב אנליסט"
        print(NameGroup)
        FolderToSave=CreateFolder(FolderToSave,"".join(x for x in NameGroup if x.isalnum()))
        FolderToSave=CreateFolder(FolderToSave,GetManpikQNow['תאריך הדוח'].max().strftime('%Y'))
        FolderToSave=CreateFolder(FolderToSave,str(int(int(GetManpikQNow['תאריך הדוח'].max().strftime('%m'))/3)))
        Listik=list(GetNameGroup.query("NameGroup==@NameGroup")["מספר תיק"].astype(int))
        ListStatus=["השגחה מיוחדת","בפיגור","מסופק","בעדכון"]
        NowManpik=GetManpikQNow.query("MisparTik in @Listik and Status in @ListStatus")["קוד מנפיק"].astype(int)
        LastQManpik=GetManpikQBefore.query("MisparTik in @Listik and Status in @ListStatus")["קוד מנפיק"].astype(int)
        #AllManpik=NowManpik.append(LastQManpik).drop_duplicates()
        AllManpik = pd.concat([NowManpik, LastQManpik]).drop_duplicates()

        PathWord="C:\\Users\\Nir.Ohad\\Desktop\\ibi-pdf\\"+"מנפיקים"+"\*.docx"
        list_of_files = glob.glob(PathWord)
        FilesToMerge=list()
        for Manpik in AllManpik:
            try:
                File=[fn for fn in list_of_files if str(Manpik) in os.path.basename(fn) ][0]
                FilesToMerge.append(File)
            except:
                pass
        PathPdf=combine_word_documents(FilesToMerge,FolderToSave)
        #Create Excel Holding File
        NameCol=["תאור מנפיק","תאור קבוצת לווים","תאור ענף"]
        df_list = {}
        for Tik in Listik:
            for Name in NameCol:
                try:
                    NameTik=GetNameGroup[GetNameGroup["מספר תיק"]==Tik]["שם מסלול"].iloc[0]
                    if NameTik=="50-60":
                        filterGetManpikQNow=GetManpikQNow.merge(GetNameGroup.rename({"מספר תיק":"MisparTik"},axis=1)[["MisparTik"]],on="MisparTik")
                    else:
                        filterGetManpikQNow=GetManpikQNow.query("MisparTik==@Tik")
                    print(3) # debug
                    if len(filterGetManpikQNow)>0:
                        print(4) # debug
                        print(filterGetManpikQNow.head()) # debug
                        AllHolding=filterGetManpikQNow.groupby([Name]).sum()[['שווי נייר']].reset_index()
                        HoldGroupByStatus=filterGetManpikQNow.groupby([Name,"Status"]).sum()[['שווי נייר']].reset_index()
                        ShoviAfik=filterGetManpikQNow['שווי נייר'].sum()
                        ShoviTik=(filterGetManpikQNow['שווי נייר']/filterGetManpikQNow['אחוז משווי תיק לפי שיערוך אחרון']).max()
                        filterGetManpikQNow["ShoviTik"]=(filterGetManpikQNow['שווי נייר']/filterGetManpikQNow['אחוז משווי תיק לפי שיערוך אחרון'])
                        ShoviTik=filterGetManpikQNow.groupby("MisparTik").max()["ShoviTik"].sum()
                        TableToExport=HoldGroupByStatus.pivot(index=Name, columns='Status', values='שווי נייר').merge(AllHolding,on=Name)
                        TableToExport["אחוז מאשראי לא מוחרג"]= TableToExport['שווי נייר']/ShoviAfik
                        TableToExport["אחוז מכלל התיק"]= TableToExport['שווי נייר']/ShoviTik
                        TableToExport=TableToExport.sort_values(by='אחוז מכלל התיק', ascending=False)
                        df_list[str(NameTik)+"-"+Name]=TableToExport
                except:
                    pass

            if NameTik =="50-60":
                filterGetManpikQNow=GetManpikQNow.merge(GetNameGroup.rename({"מספר תיק":"MisparTik"},axis=1)[["MisparTik"]],on="MisparTik")[["תאור נייר","כמות","שווי נייר","דרוג מעלות לנייר","דרוג מידרוג לנייר","תשואה ברוטו",'מח"מ מחושב',"Status"]].rename({
                    "Status":"סיווג פורום חוב"},axis=1)
            else:
                filterGetManpikQNow=GetManpikQNow.query("MisparTik==@Tik")[["תאור נייר","כמות","שווי נייר","דרוג מעלות לנייר","דרוג מידרוג לנייר","תשואה ברוטו",'מח"מ מחושב',"Status"]].rename({
                    "Status":"סיווג פורום חוב"},axis=1)
            if len(filterGetManpikQNow)>0:
                filterGetManpikQNow["אחוז מאשראי לא מוחרג"]=filterGetManpikQNow['שווי נייר']/ShoviAfik
                filterGetManpikQNow["אחוז מכלל התיק"]= filterGetManpikQNow['שווי נייר']/ShoviTik
                filterGetManpikQNow=filterGetManpikQNow.sort_values(by='אחוז מכלל התיק', ascending=False)
                df_list[str(Tik)+"-"+"כלל האחזקות"]=filterGetManpikQNow
                print(FolderToSave+"\\"+"".join(x for x in NameGroup if x.isalnum())+'.xlsx') #Debug
        writer = ExcelWriter(FolderToSave+"\\"+"".join(x for x in NameGroup if x.isalnum())+'.xlsx',engine="xlsxwriter",options={'strings_to_urls': False})
        for n in df_list:
            df=df_list[n]
            df.fillna(float(0),inplace=True)
            for row in df:
                ListStatus.append("רגיל")
                ListStatus.append("שווי נייר")

                if row in ListStatus :
                   df[row]=df[row].map(lambda x: '{:,.0f}'.format(x))
                elif row in ["אחוז מאשראי לא מוחרג","אחוז מכלל התיק"]:
                   df[row]=df[row].map(lambda x: '{:,.1f}%'.format(x*100))
                elif row in ['מח"מ מחושב',"תשואה ברוטו"]:
                   df[row]=df[row].map(lambda x: '{:,.2f}'.format(x))
            df.to_excel(writer, n[:31],startrow=1, header=False, index=False)
            workbook = writer.book
            worksheet = writer.sheets[n[:31]]
            percent_fmt = workbook.add_format({'num_format': '0.00%'})
            Dec_fmt = workbook.add_format({'num_format': '0.00'})
            total_fmt = workbook.add_format({'num_format': '#,##0'})
            if "תאור" in n:
                if len(list(df))==4:
                    worksheet.set_column('B:B', 12, total_fmt)
                    worksheet.set_column('C:D', 12, percent_fmt)
                elif len(list(df))==5:
                    worksheet.set_column('B:C', 12, total_fmt)
                    worksheet.set_column('D:E', 12, percent_fmt)
                elif len(list(df))==6:
                    worksheet.set_column('B:D', 12, total_fmt)
                    worksheet.set_column('E:F', 12, percent_fmt)
                else:
                    worksheet.set_column('B:F', 12, total_fmt)
                    worksheet.set_column('G:H', 12, percent_fmt)
            else:
                worksheet.set_column('B:C', 12, total_fmt)
                worksheet.set_column('F:G', 12, Dec_fmt)
                worksheet.set_column('I:J', 12, percent_fmt)
            column_settings = [{'header': column} for column in df.columns]
            (max_row, max_col) = df.shape
            worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings,'style': 'Table Style Medium 19'})
            format_right_to_left = workbook.add_format({'reading_order': 2})
            worksheet.right_to_left()
            for i, width in enumerate(get_col_widths(df)):
                if i==0:
                    worksheet.set_column(i, i, 30)
                else:
                    worksheet.set_column(i, i, 15)
        writer.save()
        writer.close()

        writer.handles = None
        excel_macro = win32.DispatchEx("Excel.application")
        excel_path = os.path.expanduser("DataToPDF/ExcelRun.xlsm")
        wb = excel_macro.Workbooks.Open(Filename = excel_path, ReadOnly =0)
        excel_macro.Visible = True
        wb.Sheets("Sheet1").Select() # select 2nd worksheet "Aisle_2"
        excel_macro.Range("PathFile").Value =FolderToSave+"\\"+"".join(x for x in NameGroup if x.isalnum())+'.xlsx'
        excel_macro.DisplayAlerts = False
        excel_macro.Application.Run\
             ("Module1.RunPython")
        wb.Close(False)
        excel_macro.DisplayAlerts = True
        excel_macro.Application.Quit()
        del excel_macro
        # try:
        PdfTemp = PathPdf.replace("combined_word_documents", "Temp")
        bi_pdf_path = os.path.join(os.getcwd(), "BI-PDF", "".join(x for x in NameGroup if x.isalnum()) + ".pdf")
        final_pdf_path = os.path.join(FolderToSave,
                                      "".join(x for x in NameGroup if x.isalnum()) + " דוח פורום חוב.pdf")

        # Confirm file existence
        for path, label in zip([bi_pdf_path, PathPdf, PdfTemp], ['BI-PDF', 'Combined PDF', 'Temp PDF']):
            if not os.path.exists(path):
                print(f"{label} missing: {path}")

        if all(os.path.exists(p) for p in [bi_pdf_path, PathPdf, PdfTemp]):
            print("📎 All source PDFs found — merging into final report...")
            merger = PdfMerger()
            for pdf in [bi_pdf_path, PathPdf, PdfTemp]:
                with open(pdf, 'rb') as f:
                    merger.append(f)
                    print(f"➕ Added to merge: {pdf}")
            with open(final_pdf_path, 'wb') as fout:
                merger.write(fout)
            merger.close()
            print(f"Final merged PDF created at: {final_pdf_path}")
        else:
            print("One or more source PDFs missing — skipping merge.")

    except Exception as e:
        print(f"PDF merge failed: {e}")
    # except:
    #   pass